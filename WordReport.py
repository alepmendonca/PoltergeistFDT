import re
from pathlib import Path

import win32com
import win32com.client
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm

import GeneralConfiguration
import GeneralFunctions
from GeneralFunctions import logger
from MDBReader import AIIM2003MDBReader


def _delete_paragraph(paragraph: Paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def _insert_paragraphs_before(paragraph: Paragraph, texto: str) -> Paragraph:
    texto_invertido = texto.split('\n')
    texto_invertido.reverse()
    ultimo_paragrafo = paragraph
    for t in texto_invertido:
        ultimo_paragrafo = ultimo_paragrafo.insert_paragraph_before(text=t)
    return ultimo_paragrafo


def _insert_bullets_before(paragraph: Paragraph, level_1_text: str, contents: list[str]) -> Paragraph:
    conteudos = contents
    conteudos.reverse()
    ultimo_paragrafo = paragraph
    for c in conteudos:
        ultimo_paragrafo = ultimo_paragrafo.insert_paragraph_before(
            text=f'{c}.' if c == conteudos[0] else f'{c};',
            style='Bullet List 2')
    return ultimo_paragrafo.insert_paragraph_before(text=level_1_text, style='Bullet List')


def _save_docx_as_pdf(docx_path: Path, pdf_path: Path):
    was_open = False
    try:
        word = win32com.client.GetActiveObject("Word.Application")
    except:
        word = win32com.client.Dispatch("Word.Application")
        was_open = True

    try:
        worddoc = word.Documents.Open(str(Path(docx_path).absolute()))
        worddoc.SaveAs(str(pdf_path.absolute()), FileFormat=17)
        worddoc.Close()
    finally:
        if word and was_open:
            try:
                word.Quit()
            except:
                pass


def cria_notificacao_modelo_4(cnpj: str, ie: str, razao_social: str, endereco: str,
                              corpo_notificacao: str, notification_number: str, path: Path):
    model_docx = r'resources/Notificação Modelo 4.docx'
    logger.info('Gerando notificação modelo 4...')
    docx_path = Path(str(path.parent / path.stem) + '.docx')
    doc = Document(model_docx)
    doc.core_properties.title = 'Notificação Modelo 4'

    afre_data = GeneralConfiguration.get()
    doc.tables[0].rows[0].cells[2].paragraphs[3].runs[0].text = f'NOTIFICAÇÃO {notification_number}'
    doc.tables[0].rows[2].cells[0].paragraphs[0].runs[0].text = afre_data.drt_sigla
    doc.tables[0].rows[2].cells[1].paragraphs[0].runs[0].text = f'NF-{afre_data.nucleo_fiscal()}'
    doc.tables[0].rows[2].cells[3].paragraphs[0].runs[0].text = f'EQUIPE {afre_data.equipe_fiscal}'
    with AIIM2003MDBReader() as aiim2003:
        funcional = int(re.sub(r'\D', '', afre_data.funcional))
        endereco_drt, telefone_drt = aiim2003.get_afr_drt_address_and_phone(funcional)
    doc.tables[0].rows[4].cells[0].paragraphs[0].runs[0].text = endereco_drt
    doc.tables[0].rows[4].cells[4].paragraphs[0].runs[0].text = telefone_drt
    doc.tables[1].rows[2].cells[0].paragraphs[0].runs[0].text = f'{razao_social}\n{endereco}'
    doc.tables[1].rows[2].cells[1].paragraphs[0].runs[0].text = ie
    doc.tables[1].rows[4].cells[1].paragraphs[0].runs[0].text = {cnpj}
    # faz conversão de HTML
    celula_corpo = doc.tables[2].rows[5].cells[0]
    paragrafo_corpo = celula_corpo.paragraphs[0]
    paragrafo_corpo.runs[0].text = ''
    texto = corpo_notificacao.replace('<br>', '\n').replace('<p>', '\n').replace('</p>', '').\
        replace('<ul>', '').replace('</ul>', '').replace('</li>', '\n')
    texto = '\n' + texto + ''.ljust(20 - texto.count('\n'), '\n')
    linhas = texto.splitlines()
    for idx in range(0, len(linhas)):
        paragrafo = celula_corpo.add_paragraph('') if idx > 0 else paragrafo_corpo
        paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        linha = linhas[idx]
        while linha.find('<') >= 0:
            paragrafo.add_run(linha[:linha.index('<')])
            comando = linha[linha.index('<'):linha.index('>')+1]
            linha = linha[linha.index('>')+1:]
            match comando:
                case '<b>':
                    paragrafo.add_run(linha[:linha.index('</b>')]).bold = True
                    linha = linha[linha.index('</b>')+4:]
                case '<i>':
                    paragrafo.add_run(linha[:linha.index('</i>')]).italic = True
                    linha = linha[linha.index('</i>') + 4:]
                case '<li>':
                    paragrafo = celula_corpo.add_paragraph('')
                    paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    paragrafo.style = 'Bullet List'
                case _:
                    raise Exception(f'Não sei tratar o HTML corretamente: {corpo_notificacao}')
        if linha:
            paragrafo.add_run(linha)
    doc.tables[2].rows[10].cells[0].paragraphs[0].runs[0].text = f'{afre_data.drt_sigla}-' \
                                                                 f'NF-{afre_data.nucleo_fiscal()}-' \
                                                                 f'EQ-{afre_data.equipe_fiscal},' \
                                                                 f' ______/______/_______ '
    doc.tables[2].rows[-1].cells[0].paragraphs[1].runs[0].text = afre_data.nome
    doc.tables[2].rows[-1].cells[0].paragraphs[2].runs[0].text = f'AFRE - IF {afre_data.funcional}'
    doc.save(str(docx_path.absolute()))
    _save_docx_as_pdf(docx_path, path)

def cria_recibo_entrega_arquivos_digitais(cnpj: str, ie: str, razao_social: str, osf: str,
                                          hashes: list[dict], path: Path):
    docx_path = Path(str(path.parent / path.stem) + '.docx')
    try:
        model_docx = r'resources/Recibo de Entrega de Arquivos Digitais.docx'
        logger.info('Gerando PDF do recibo de entrega de arquivos digitais...')
        doc = Document(model_docx)
        doc.core_properties.title = 'Recibo de Entrega de Arquivos Digitais'
        doc.tables[0].rows[0].cells[1].text = cnpj
        doc.tables[0].rows[0].cells[4].text = ie
        doc.tables[0].rows[1].cells[2].text = razao_social
        doc.tables[1].rows[0].cells[0].paragraphs[0].text = \
            doc.tables[1].rows[0].cells[0].paragraphs[0].text.replace('<osf>', osf)

        for dados in hashes:
            doc.add_paragraph()
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run(f'{dados["notification"]}')
            run.font.bold = True
            run.font.size = Pt(9)

            table = doc.add_table(rows=len(dados['files']) + 1, cols=4)
            table.autofit = False
            table.allow_autofit = False
            widths = (Cm(9.75), Cm(2), Cm(5.25), Cm(10.5))
            for row in table.rows:
                for idx, width in enumerate(widths):
                    row.cells[idx].width = width
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Grid Table 4 Accent 3'
            hdr_cells = table.rows[0].cells
            paragraph = hdr_cells[0].paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run('Nome do Arquivo')
            run.font.bold = True
            run.font.size = Pt(9)
            paragraph = hdr_cells[1].paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run('Tamanho')
            run.font.bold = True
            run.font.size = Pt(9)
            paragraph = hdr_cells[2].paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run('Assinatura MD5')
            run.font.bold = True
            run.font.size = Pt(9)
            paragraph = hdr_cells[3].paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run('Assinatura SHA-256')
            run.font.bold = True
            run.font.size = Pt(9)
            for i in range(1, len(dados['files']) + 1):
                table.rows[i].cells[0].text = dados['files'][i - 1]['file']
                table.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(8)
                table.rows[i].cells[1].text = dados['files'][i - 1]['size']
                table.rows[i].cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                table.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(8)
                table.rows[i].cells[2].text = dados['files'][i - 1]['md5']
                table.rows[i].cells[2].paragraphs[0].runs[0].font.size = Pt(8)
                table.rows[i].cells[3].text = dados['files'][i - 1]['sha256']
                table.rows[i].cells[3].paragraphs[0].runs[0].font.size = Pt(8)

        doc.add_paragraph('')
        paragraph = doc.add_paragraph(GeneralConfiguration.get().nome)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph = doc.add_paragraph('AUDITOR FISCAL DA RECEITA ESTADUAL')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph = doc.add_paragraph(f'IF {GeneralConfiguration.get().funcional}')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.save(str(docx_path.absolute()))
        _save_docx_as_pdf(docx_path, path)
    finally:
        docx_path.unlink(missing_ok=True)


def cria_capa_para_anexo(title: str, path: Path):
    docx_path = Path(str(path.parent / path.stem) + '.docx')
    logger.info('Gerando capa do Anexo...')
    try:
        doc = Document()
        for i in range(1, 10):
            doc.add_paragraph('')
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run(title)
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(72)
        doc.save(str(docx_path.absolute()))
        _save_docx_as_pdf(docx_path, path)
    finally:
        docx_path.unlink(missing_ok=True)


class WordReport:
    PROVAS = 1000
    PROVAS_GERAIS_BULLET = 100
    titulos = {
        1: "DAS IRREGULARIDADES NO PAGAMENTO DO IMPOSTO",
        2: "DAS IRREGULARIDADES NO CRÉDITO DO IMPOSTO",
        3: "DAS IRREGULARIDADES NA DOCUMENTAÇÃO FISCAL NA ENTREGA, REMESSA, TRANSPORTE, RECEBIMENTO, "
           "ESTOCAGEM, DEPÓSITO OU PRESTAÇÃO",
        4: "DAS IRREGULARIDADES NOS DOCUMENTOS FISCAIS",
        5: "DAS IRREGULARIDADES NOS LIVROS FISCAIS",
        6: "DAS IRREGULARIDADES NO CADASTRO DO CONTRIBUINTE",
        7: "DAS IRREGULARIDADES NA APRESENTAÇÃO DE INFORMAÇÃO ECONÔMICO-FISCAL E GUIA DE RECOLHIMENTO DE IMPOSTO",
        8: "DAS IRREGULARIDADES EM SISTEMA DE PROCESSAMENTO DE DADOS OU EQUIPAMENTOS",
        9: "DAS IRREGULARIDADES NA INTERVENÇÃO TÉCNICA EM EMISSOR DE CUPOM FISCAL",
        10: "DAS IRREGULARIDADES NO DESENVOLVIMENTO DE SOFTWARE APLICATIVO PARA EMISSOR DE CUPOM FISCAL",
        11: "DAS DEMAIS IRREGULARIDADES",
        PROVAS: "DA CONCLUSÃO DO TRABALHO FISCAL"
    }

    def __init__(self, audit):
        self._audit = audit
        self.report: Document = None
        audit.aiim_path().mkdir(exist_ok=True)
        self.report_path: Path = audit.aiim_path() / 'Relatório Circunstanciado.docx'
        if not self.report_path.is_file():
            try:
                self.report = Document(Path('resources/Relatório Circunstanciado.docx'))
                self._initialize_report()
                self.save_report()
            except Exception as e:
                logger.exception('Ocorreu uma falha ao copiar a planilha template para a pasta da fiscalização')
                try:
                    self.report_path.unlink(missing_ok=True)
                except Exception:
                    pass
                raise e
        else:
            self.report: Document = Document(self.report_path)

    def _initialize_report(self):
        self.report.sections[0].header.tables[0].rows[0].cells[0].paragraphs[2].add_run(
            f'{GeneralConfiguration.get().drt_nome} - {GeneralConfiguration.get().drt_sigla}\n'
            f'NÚCLEO DE FISCALIZAÇÃO {GeneralConfiguration.get().nucleo_fiscal()} '
            f'- EQUIPE {GeneralConfiguration.get().equipe_fiscal}')
        for paragraph in self.report.paragraphs:
            paragraph.text = paragraph.text.replace('<osf>', self._audit.osf)
            paragraph.text = paragraph.text.replace('<aiim>', self._audit.aiim_number)
            paragraph.text = paragraph.text.replace('<afre>', GeneralConfiguration.get().nome)
            paragraph.text = paragraph.text.replace('<if>', GeneralConfiguration.get().funcional)
            paragraph.text = paragraph.text.replace('<delegacia-sigla>', GeneralConfiguration.get().drt_sigla)
            paragraph.text = paragraph.text.replace('<nf>', f'{GeneralConfiguration.get().nucleo_fiscal()}')
            paragraph.text = paragraph.text.replace('<equipe>', f'{GeneralConfiguration.get().equipe_fiscal}')
            if paragraph.text.startswith('RELATÓRIO'):
                for run in paragraph.runs:
                    run.bold = True
                    run.underline = True
            if paragraph.text.startswith('(Anexo'):
                for run in paragraph.runs:
                    run.bold = True
            if paragraph.text.startswith('FISCALIZADA:'):
                paragraph.clear()
                paragraph.add_run(f'FISCALIZADA: {self._audit.empresa}\nIE: {self._audit.ie}\n'
                                  f'CNPJ: {self._audit.cnpj}\n'
                                  f'LOCALIDADE: {self._audit.endereco_completo()}').bold = True
            if paragraph.text.startswith('O contribuinte está enquadrado'):
                paragraph.clear()
                paragraph.add_run(self._audit.periodos_da_fiscalizacao_descricao())

    def insere_item(self, item: int, inciso: int, texto: str, tem_provas: bool):
        if item <= 0:
            return
        titulos_existentes = self.titulos_inseridos()
        texto = f'No item {item}, {texto}'
        if tem_provas:
            texto += f'\nOs demonstrativos que comprovam especificamente a infringência deste item ' \
                     f'do auto de infração encontram-se no Anexo do Item {item}.'
        texto += '\n'
        if inciso in titulos_existentes:
            # caso já exista o título, coloca dentro dele
            idx_titulo = titulos_existentes[inciso]
            idx_prox_titulo = titulos_existentes[
                list(titulos_existentes.keys())[list(titulos_existentes.keys()).index(inciso) + 1]]
            for idx in range(idx_titulo + 1, idx_prox_titulo + 1):
                texto_item = re.search(r'^No item (\d+)', self.report.paragraphs[idx].text)
                if texto_item:
                    item_existente = int(texto_item.group(1))
                    if item_existente > item:
                        _insert_paragraphs_before(self.report.paragraphs[idx], texto)
                        return
            # não achou item maior, então insere no final
            _insert_paragraphs_before(self.report.paragraphs[idx_prox_titulo - 1], texto)
        else:
            # não existe o título, então cria ele e bota o texto em seguida
            for alinea_existente in titulos_existentes.keys():
                if alinea_existente > inciso:
                    ultimo_paragrafo = _insert_paragraphs_before(
                        self.report.paragraphs[titulos_existentes[alinea_existente]], texto)
                    ultimo_paragrafo.insert_paragraph_before(text=f'- {self.titulos[inciso]}', style='Heading 1')
                    break

    def titulos_inseridos(self) -> dict:
        retorno = {}
        for i in range(len(self.report.paragraphs)):
            if self.report.paragraphs[i].style.name == 'Heading 1':
                if self.report.paragraphs[i].text[2:].strip() in self.titulos.values():
                    retorno[list(self.titulos.keys())[
                        list(self.titulos.values()).index(self.report.paragraphs[i].text[2:].strip())]] = i
        return retorno

    def insere_anexo(self, item: int, conteudos: list[str]):
        if item <= 0:
            return
        ultimo_paragrafo: Paragraph = None
        for anexo_existente, idx in self.anexos_inseridos().items():
            if anexo_existente > item:
                ultimo_paragrafo = self.report.paragraphs[idx]
                break
        if not ultimo_paragrafo:
            raise Exception('Não encontrei lugar certo para colocar Anexo! Verifique o template do relatório!')
        _insert_bullets_before(ultimo_paragrafo, f'Anexo do Item {item}, contendo:', conteudos)

    def atualiza_provas_gerais(self, conteudos: list[str]):
        bullet_provas_gerais = self.anexos_inseridos()[self.PROVAS_GERAIS_BULLET]
        _insert_bullets_before(self.report.paragraphs[bullet_provas_gerais], 'Provas Gerais, contendo:', conteudos)
        a_remover = []
        for idx in range(bullet_provas_gerais + len(conteudos) + 1, len(self.report.paragraphs)):
            if self.report.paragraphs[idx].text == '':
                break
            a_remover.append(idx)
        a_remover.reverse()
        for b in a_remover:
            _delete_paragraph(self.report.paragraphs[b])

    def anexos_inseridos(self) -> dict[int, int]:
        retorno = {}
        for i in range(self.titulos_inseridos()[self.PROVAS], len(self.report.paragraphs)):
            texto_anexo = re.search(r'^Anexo do Item (\d+)', self.report.paragraphs[i].text)
            if texto_anexo:
                anexo_existente = int(texto_anexo.group(1))
                retorno[anexo_existente] = i
            elif self.report.paragraphs[i].text.startswith('Provas Gerais'):
                retorno[self.PROVAS_GERAIS_BULLET] = i
        return retorno

    def remove_item(self, item: int):
        if item <= 0:
            return
        pos = -1
        for i in range(len(self.report.paragraphs)):
            if pos == -1 and self.report.paragraphs[i].text.startswith(f'No item {item},'):
                pos = i
                _delete_paragraph(self.report.paragraphs[pos])
                while not self.report.paragraphs[i].text.startswith(f'No item') \
                        and self.report.paragraphs[i].style.name != 'Heading 1':
                    _delete_paragraph(self.report.paragraphs[i])
                break
        if pos > -1:
            # caso era o último item dentro do título, apaga o título também
            if self.report.paragraphs[pos - 1].style.name == 'Heading 1' and \
                    self.report.paragraphs[pos].style.name == 'Heading 1':
                _delete_paragraph(self.report.paragraphs[pos - 1])

    def remove_anexo(self, item: int):
        if item <= 0:
            return
        for i in range(len(self.report.paragraphs)):
            if self.report.paragraphs[i].text.startswith(f'Anexo do Item {item},'):
                _delete_paragraph(self.report.paragraphs[i])
                while not self.report.paragraphs[i].text.startswith(f'Anexo') \
                        and not self.report.paragraphs[i].text.startswith(f'Provas Gerais'):
                    _delete_paragraph(self.report.paragraphs[i])
                break

    def save_report(self):
        self.report.save(self.report_path)
